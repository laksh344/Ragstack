"""Multi-format document parser.

Parses PDF, DOCX, CSV, TXT files into structured ParsedDocument objects.
Each page/section preserves metadata for downstream chunking and attribution.

Reference: RAGFlow's deepdoc/parser/ splits by format similarly,
but we use PyMuPDF + python-docx instead of their custom C++ parsers.
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

import pandas as pd
import structlog

if TYPE_CHECKING:
    from backend.models.document import ParsedDocument

logger = structlog.get_logger()


def parse_document(file_path: str | Path) -> ParsedDocument:
    """Route to the correct parser based on file extension."""
    from backend.models.document import FileType

    path = Path(file_path)
    file_type = FileType.from_extension(path.suffix)

    parser_map = {
        FileType.PDF: _parse_pdf,
        FileType.DOCX: _parse_docx,
        FileType.CSV: _parse_csv,
        FileType.TXT: _parse_txt,
        FileType.XLSX: _parse_csv,  # reuse CSV logic via pandas
    }

    parser_fn = parser_map.get(file_type)
    if not parser_fn:
        raise ValueError(f"Unsupported file type: {path.suffix}")

    logger.info("parser.start", file=path.name, file_type=file_type.value)
    doc = parser_fn(path)
    doc.source_file = path.name
    doc.file_type = file_type
    logger.info(
        "parser.complete",
        file=path.name,
        pages=doc.total_pages,
        chars=sum(len(p.content) for p in doc.pages),
    )
    return doc


def _parse_pdf(path: Path) -> ParsedDocument:
    """Parse PDF using PyMuPDF with table detection.

    PyMuPDF extracts text page-by-page and can detect tables natively.
    Pages with tables or images are flagged for optional vision processing.
    """
    import fitz  # pymupdf

    from backend.models.document import ParsedDocument, ParsedPage

    doc = fitz.open(str(path))
    pages = []

    for page_num, page in enumerate(doc, start=1):
        # Extract text
        text = page.get_text("text").strip()

        # Detect tables using PyMuPDF's built-in table finder
        tables = page.find_tables()
        table_data = []
        has_tables = len(tables.tables) > 0

        for table in tables:
            try:
                df = table.to_pandas()
                # Convert table to markdown for LLM consumption
                md_table = df.to_markdown(index=False)
                table_data.append(md_table)
            except Exception:
                # Fallback: extract as CSV-like text
                for row in table.extract():
                    row_text = " | ".join(str(cell) if cell else "" for cell in row)
                    table_data.append(row_text)

        # Detect images
        image_list = page.get_images(full=True)
        has_images = len(image_list) > 0

        # Combine text with table markdown
        full_content = text
        if table_data:
            full_content += "\n\n--- Tables ---\n" + "\n\n".join(table_data)

        pages.append(
            ParsedPage(
                page_number=page_num,
                content=full_content,
                has_tables=has_tables,
                has_images=has_images,
                table_data=table_data,
            )
        )

    # Extract document metadata
    pdf_metadata = doc.metadata or {}
    title = pdf_metadata.get("title", "") or path.stem

    doc.close()

    return ParsedDocument(
        source_file=path.name,
        title=title,
        total_pages=len(pages),
        pages=pages,
        metadata={
            "author": pdf_metadata.get("author", ""),
            "subject": pdf_metadata.get("subject", ""),
            "creator": pdf_metadata.get("creator", ""),
            "creation_date": pdf_metadata.get("creationDate", ""),
        },
    )


def _parse_docx(path: Path) -> ParsedDocument:
    """Parse DOCX using python-docx.

    Extracts paragraphs and tables, grouping content into logical sections.
    Tables are converted to markdown format.
    """
    from docx import Document as DocxDocument
    from docx.table import Table

    from backend.models.document import ParsedDocument, ParsedPage

    doc = DocxDocument(str(path))
    pages = []
    current_content: list[str] = []
    page_num = 1

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Paragraph
            from docx.text.paragraph import Paragraph

            para = Paragraph(element, doc)
            text = para.text.strip()
            if text:
                current_content.append(text)

            # Check for page break
            for run in para.runs:
                if run._element.xml.find("w:br") != -1 and 'type="page"' in run._element.xml:
                    if current_content:
                        pages.append(
                            ParsedPage(
                                page_number=page_num,
                                content="\n".join(current_content),
                            )
                        )
                        current_content = []
                        page_num += 1

        elif tag == "tbl":
            # Table
            table = Table(element, doc)
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)

            if rows:
                # Convert to markdown table
                header = " | ".join(rows[0])
                separator = " | ".join("---" for _ in rows[0])
                body = "\n".join(" | ".join(row) for row in rows[1:])
                md_table = f"{header}\n{separator}\n{body}"
                current_content.append(f"\n{md_table}\n")

    # Flush remaining content
    if current_content:
        pages.append(
            ParsedPage(page_number=page_num, content="\n".join(current_content))
        )

    title = doc.core_properties.title or path.stem

    return ParsedDocument(
        source_file=path.name,
        title=title,
        total_pages=len(pages),
        pages=pages,
        metadata={
            "author": doc.core_properties.author or "",
            "created": str(doc.core_properties.created or ""),
            "modified": str(doc.core_properties.modified or ""),
        },
    )


def _parse_csv(path: Path) -> ParsedDocument:
    """Parse CSV/XLSX into a single-page document with structured text.

    Converts tabular data into a text representation that preserves
    column names and relationships for effective RAG retrieval.
    """
    from backend.models.document import ParsedDocument, ParsedPage

    if path.suffix.lower() == ".xlsx":
        df = pd.read_excel(path)
    else:
        df = pd.read_csv(path)

    # Create structured text representation
    lines = [
        f"Dataset: {path.stem}",
        f"Columns: {', '.join(df.columns.tolist())}",
        f"Rows: {len(df)}",
        "",
        "--- Data Preview (first 50 rows) ---",
        df.head(50).to_markdown(index=False),
        "",
        "--- Column Statistics ---",
    ]

    # Add statistics for numeric columns
    numeric_cols = df.select_dtypes(include=["number"]).columns
    if len(numeric_cols) > 0:
        lines.append(df[numeric_cols].describe().to_markdown())

    # Add unique value counts for categorical columns
    cat_cols = df.select_dtypes(include=["object"]).columns
    for col in cat_cols[:5]:  # limit to first 5 categorical columns
        unique_count = df[col].nunique()
        top_values = df[col].value_counts().head(5).to_dict()
        lines.append(f"\n{col}: {unique_count} unique values. Top: {top_values}")

    content = "\n".join(lines)

    return ParsedDocument(
        source_file=path.name,
        title=path.stem,
        total_pages=1,
        pages=[ParsedPage(page_number=1, content=content, has_tables=True)],
        metadata={
            "rows": len(df),
            "columns": len(df.columns),
            "column_names": df.columns.tolist(),
        },
    )


def _parse_txt(path: Path) -> ParsedDocument:
    """Parse plain text file, splitting into pages by line count."""
    from backend.models.document import ParsedDocument, ParsedPage

    text = path.read_text(encoding="utf-8", errors="replace")
    lines = text.split("\n")

    # Split into ~50-line pages for manageable chunks
    page_size = 50
    pages = []
    for i in range(0, len(lines), page_size):
        page_lines = lines[i : i + page_size]
        content = "\n".join(page_lines).strip()
        if content:
            pages.append(
                ParsedPage(page_number=(i // page_size) + 1, content=content)
            )

    return ParsedDocument(
        source_file=path.name,
        title=path.stem,
        total_pages=len(pages),
        pages=pages,
        metadata={"total_lines": len(lines), "total_chars": len(text)},
    )
